import random
import os
from docx import Document
from docx.shared import Inches
from PIL import Image


MAX_IMAGE_WIDTH_INCHES = 3
NUMBER_COL_WIDTH = Inches(0.8)
SCHEMES_FOLDER = 'schemes'
OUTPUT_DOCX = 'generated_schemes.docx'


def add_schemes_to_word(scheme_type, save_path):
    doc = Document()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.allow_autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер варианта'
    hdr_cells[1].text = 'Схема'
    hdr_cells[2].text = 'Номиналы'
    hdr_cells[0].width = NUMBER_COL_WIDTH
    hdr_cells[1].width = MAX_IMAGE_WIDTH_INCHES
    hdr_cells[2].width = Inches(2.5)

    for i in range(1, 31):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[0].width = NUMBER_COL_WIDTH
        row_cells[1].width = MAX_IMAGE_WIDTH_INCHES
        row_cells[2].width = Inches(2.5)

        image_path = os.path.join(save_path, SCHEMES_FOLDER, f'scheme_{i}.png')

        if os.path.exists(image_path):
            with Image.open(image_path) as img:
                width_px, height_px = img.size
                dpi = img.info.get('dpi', (300, 300))[0]
                width_in = width_px / dpi
                scale = min(1.0, MAX_IMAGE_WIDTH_INCHES / width_in)
                display_width = Inches(width_in * scale)
                row_cells[1].paragraphs[0].add_run().add_picture(image_path, width=display_width)

                metadata = img.info
                switch_info = metadata.get('switch_info', "no")
        else:
            row_cells[1].text = 'Изображение не найдено'

        descriptions = []

        voltages = []
        currents = []
        resistors = []

        for element_name, element_value in metadata.items():
            if element_name.startswith('V'):
                descriptions.append(f"{element_name}={element_value}")
                voltages.append(element_name)
            elif element_name.startswith('I'):
                descriptions.append(f"{element_name}={element_value.replace('.', ',')}")
                currents.append(element_name)
            elif element_name.startswith('R'):
                descriptions.append(f"{element_name}={element_value}")
                resistors.append(element_name)
            elif element_name.startswith('C'):
                descriptions.append(f"{element_name}={element_value.replace('.', ',')}")
            elif element_name.startswith('L'):
                descriptions.append(f"{element_name}={element_value}")

        if scheme_type == "alternating_current" or scheme_type == "transient_processes":
            frequency = random.randint(1, 20)
            descriptions.append(f"f={frequency} кГц")

        if scheme_type == "coupling_coefficient":
            if (voltages or currents) and resistors:
                is_voltage = random.choice([True, False])
                if is_voltage and voltages:
                    v_index = random.randint(1, len(voltages))
                    source_label = f"V{v_index}"
                elif not is_voltage and currents:
                    c_index = random.randint(1, len(currents))
                    source_label = f"I{c_index}"
                else:
                    source_label = f"V1" if voltages else f"I1"

            r_index = random.randint(1, len(resistors))
            task = f"Рассчитать коэффициент связи между {source_label} и R{r_index}"
            descriptions.append("")
            descriptions.append(task)

        if switch_info != "no":
            if switch_info == "opening":
                descriptions.append("Ключ замыкается")
            elif switch_info == "closing":
                descriptions.append("Ключ размыкается")

        row_cells[2].text = '\n'.join(descriptions)

    doc.save(f'{save_path}/{OUTPUT_DOCX}')
    print(f'Файл сохранён: {OUTPUT_DOCX}')
