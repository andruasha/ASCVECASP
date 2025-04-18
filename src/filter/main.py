import matplotlib.pyplot as plt
import numpy as np
import random
import os

from src.common.draw_functions import draw_resistor, draw_capacitor, draw_inductor

from docx import Document
from docx.shared import Inches
from PIL import Image

from conf.config import SCALE

SCHEMES_FOLDER = 'schemes'
MAX_IMAGE_WIDTH_INCHES = 3
NUMBER_COL_WIDTH = Inches(0.8)
OUTPUT_DOCX = 'generated_schemes.docx'


def generate_filter_scheme_topology(scheme_type):
    all_nodes = {}
    scheme_nodes = {}
    quadripole_nodes = {}
    scheme_layout = {}
    if scheme_type == 'G':
        all_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node2': {'x': SCALE, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node4': {'x': 0, 'y': 0},
            'node5': {'x': SCALE, 'y': 0},
            'node6': {'x': 3 * SCALE / 2, 'y': 0}
        }

        scheme_nodes = {
            'node2': {'x': SCALE, 'y': SCALE},
            'node5': {'x': SCALE, 'y': 0}
        }

        quadripole_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node4': {'x': 0, 'y': 0},
            'node6': {'x': 3 * SCALE / 2, 'y': 0}
        }

        scheme_layout = {
            'node1->node2': [
                {
                    'connection_coords': [
                        all_nodes['node1'],
                        all_nodes['node2']
                    ],
                    'elements': []
                }
            ],
            'node2->node3': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node3']
                    ]
                }
            ],
            'node4->node5': [
                {
                    'connection_coords': [
                        all_nodes['node4'],
                        all_nodes['node5']
                    ]
                }
            ],
            'node5->node6': [
                {
                    'connection_coords': [
                        all_nodes['node5'],
                        all_nodes['node6']
                    ]
                }
            ],
            'node2->node5': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node5']
                    ],
                    'elements': []
                }
            ]
        }

    elif scheme_type == 'P':
        all_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node2': {'x': SCALE / 2, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node4': {'x': 2 * SCALE, 'y': SCALE},
            'node5': {'x': 0, 'y': 0},
            'node6': {'x': SCALE / 2, 'y': 0},
            'node7': {'x': 3 * SCALE / 2, 'y': 0},
            'node8': {'x': 2 * SCALE, 'y': 0}
        }

        scheme_nodes = {
            'node2': {'x': SCALE / 2, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node6': {'x': SCALE / 2, 'y': 0},
            'node7': {'x': 3 * SCALE / 2, 'y': 0}
        }

        quadripole_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node4': {'x': 2 * SCALE, 'y': SCALE},
            'node5': {'x': 0, 'y': 0},
            'node8': {'x': 2 * SCALE, 'y': 0}
        }

        scheme_layout = {
            'node1->node2': [
                {
                    'connection_coords': [
                        all_nodes['node1'],
                        all_nodes['node2']
                    ]
                }
            ],
            'node2->node3': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node3']
                    ],
                    'elements': []
                }
            ],
            'node3->node4': [
                {
                    'connection_coords': [
                        all_nodes['node3'],
                        all_nodes['node4']
                    ]
                }
            ],
            'node5->node6': [
                {
                    'connection_coords': [
                        all_nodes['node5'],
                        all_nodes['node6']
                    ]
                }
            ],
            'node6->node7': [
                {
                    'connection_coords': [
                        all_nodes['node6'],
                        all_nodes['node7']
                    ]
                }
            ],
            'node7->node8': [
                {
                    'connection_coords': [
                        all_nodes['node7'],
                        all_nodes['node8']
                    ]
                }
            ],
            'node2->node6': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node6']
                    ],
                    'elements': []
                }
            ],
            'node3->node7': [
                {
                    'connection_coords': [
                        all_nodes['node3'],
                        all_nodes['node7']
                    ],
                    'elements': []
                }
            ]
        }

    elif scheme_type == 'T':
        all_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node2': {'x': SCALE, 'y': SCALE},
            'node3': {'x': 2 * SCALE, 'y': SCALE},
            'node4': {'x': 0, 'y': 0},
            'node5': {'x': SCALE, 'y': 0},
            'node6': {'x': 2 * SCALE, 'y': 0}
        }

        scheme_nodes = {
            'node2': {'x': SCALE, 'y': SCALE},
            'node5': {'x': SCALE, 'y': 0}
        }

        quadripole_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node3': {'x': 2 * SCALE, 'y': SCALE},
            'node4': {'x': 0, 'y': 0},
            'node6': {'x': 2 * SCALE, 'y': 0}
        }

        scheme_layout = {
            'node1->node2': [
                {
                    'connection_coords': [
                        all_nodes['node1'],
                        all_nodes['node2']
                    ],
                    'elements': []
                }
            ],
            'node2->node3': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node3']
                    ],
                    'elements': []
                }
            ],
            'node4->node5': [
                {
                    'connection_coords': [
                        all_nodes['node4'],
                        all_nodes['node5']
                    ]
                }
            ],
            'node5->node6': [
                {
                    'connection_coords': [
                        all_nodes['node5'],
                        all_nodes['node6']
                    ]
                }
            ],
            'node2->node5': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node5']
                    ],
                    'elements': []
                }
            ],
        }

    elif scheme_type == 'T_bridge':
        all_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node2': {'x': SCALE / 2, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node4': {'x': 2 * SCALE, 'y': SCALE},
            'node5': {'x': 0, 'y': 0},
            'node6': {'x': SCALE / 2, 'y': 0},
            'node7': {'x': 3 * SCALE / 2, 'y': 0},
            'node8': {'x': 2 * SCALE, 'y': 0}
        }

        scheme_nodes = {
            'node2': {'x': SCALE / 2, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node6': {'x': SCALE / 2, 'y': 0},
            'node7': {'x': 3 * SCALE / 2, 'y': 0}
        }

        quadripole_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node4': {'x': 2 * SCALE, 'y': SCALE},
            'node5': {'x': 0, 'y': 0},
            'node8': {'x': 2 * SCALE, 'y': 0}
        }

        scheme_layout = {
            'node1->node2': [
                {
                    'connection_coords': [
                        all_nodes['node1'],
                        all_nodes['node2']
                    ]
                }
            ],
            'node2->node3': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node3']
                    ],
                    'elements': []
                }
            ],
            'node3->node4': [
                {
                    'connection_coords': [
                        all_nodes['node3'],
                        all_nodes['node4']
                    ]
                }
            ],
            'node5->node6': [
                {
                    'connection_coords': [
                        all_nodes['node5'],
                        all_nodes['node6']
                    ]
                }
            ],
            'node6->node7': [
                {
                    'connection_coords': [
                        all_nodes['node6'],
                        all_nodes['node7']
                    ],
                    'elements': []
                }
            ],
            'node7->node8': [
                {
                    'connection_coords': [
                        all_nodes['node7'],
                        all_nodes['node8']
                    ]
                }
            ],
            'node2->node7': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node7']
                    ],
                    'elements': []
                }
            ],
            'node3->node6': [
                {
                    'connection_coords': [
                        all_nodes['node3'],
                        all_nodes['node6']
                    ],
                    'elements': []
                }
            ]
        }

    elif scheme_type == 'T_back_coupling':
        all_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node2': {'x': SCALE / 2, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node4': {'x': 2 * SCALE, 'y': 3 * SCALE / 2},
            'node5': {'x': 3 * SCALE, 'y': SCALE},
            'node6': {'x': 7 * SCALE / 2, 'y': SCALE},
            'node7': {'x': 0, 'y': 0},
            'node8': {'x': 3 * SCALE / 2, 'y': 0},
            'node9': {'x': 2 * SCALE, 'y': 0},
            'node10': {'x': 7 * SCALE / 2, 'y': 0}
        }

        scheme_nodes = {
            'node2': {'x': SCALE / 2, 'y': SCALE},
            'node3': {'x': 3 * SCALE / 2, 'y': SCALE},
            'node4': {'x': 2 * SCALE, 'y': 3 * SCALE / 2},
            'node5': {'x': 3 * SCALE, 'y': SCALE},
            'node8': {'x': 3 * SCALE / 2, 'y': 0},
            'node9': {'x': 2 * SCALE, 'y': 0},
        }

        quadripole_nodes = {
            'node1': {'x': 0, 'y': SCALE},
            'node6': {'x': 7 * SCALE / 2, 'y': SCALE},
            'node7': {'x': 0, 'y': 0},
            'node10': {'x': 7 * SCALE / 2, 'y': 0}
        }

        scheme_layout = {
            'node1->node2': [
                {
                    'connection_coords': [
                        all_nodes['node1'],
                        all_nodes['node2']
                    ]
                }
            ],
            'node2->node3': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        all_nodes['node3']
                    ],
                    'elements': []
                }
            ],
            'node3->node5': [
                {
                    'connection_coords': [
                        all_nodes['node3'],
                        all_nodes['node5']
                    ],
                    'elements': []
                }
            ],
            'node5->node6': [
                {
                    'connection_coords': [
                        all_nodes['node5'],
                        all_nodes['node6']
                    ]
                }
            ],
            'node2->node4': [
                {
                    'connection_coords': [
                        all_nodes['node2'],
                        {'x': all_nodes['node2']['x'], 'y': all_nodes['node4']},
                        all_nodes['node4']
                    ],
                    'elements': []
                }
            ],
            'node4->node5': [
                {
                    'connection_coords': [
                        all_nodes['node4'],
                        {'x': all_nodes['node5']['x'], 'y': all_nodes['node4']},
                        all_nodes['node4']
                    ],
                    'elements': []
                }
            ],
            'node7->node8': [
                {
                    'connection_coords': [
                        all_nodes['node7'],
                        all_nodes['node8']
                    ]
                }
            ],
            'node8->node9': [
                {
                    'connection_coords': [
                        all_nodes['node8'],
                        all_nodes['node9']
                    ]
                }
            ],
            'node9->node10': [
                {
                    'connection_coords': [
                        all_nodes['node9'],
                        all_nodes['node10']
                    ]
                }
            ]
        }

    return scheme_nodes, quadripole_nodes, scheme_layout


def visualise_filter_scheme(scheme_nodes, quadripole_nodes, scheme_layout):
    plt.figure(figsize=(5, 5))
    plt.grid(True, which='both', linestyle='--', linewidth=0.5, color='gray')
    plt.xticks(range(-13, 14, 1))
    plt.yticks(range(-13, 14, 1))
    plt.xlim(-13, 13)
    plt.ylim(-13, 13)
    plt.axis('equal')

    for node, coords in scheme_nodes.items():
        plt.plot(coords['x'], coords['y'], 'ko', markersize=3)
    for node, coords in quadripole_nodes.items():
        plt.plot(coords['x'], coords['y'], 'ro', markersize=3, zorder=11)

    resistor_idx = 1
    capacitor_idx = 1
    inductor_idx = 1

    for connections in scheme_layout.values():
        for segment in connections:
            coords = segment['connection_coords']
            for i in range(len(coords) - 1):
                start = coords[i]
                end = coords[i + 1]
                plt.plot([start['x'], end['x']], [start['y'], end['y']], 'k-', linewidth=1)

            if 'elements' not in segment:
                continue

            elements = segment['elements']
            num_elements = len(elements)

            if num_elements == 0:
                continue

            lines = []
            for i in range(len(coords) - 1):
                lines.append((coords[i], coords[i + 1]))

            num_lines = len(lines)
            elements_per_line = num_elements // num_lines
            extra = num_elements % num_lines

            element_iter = iter(elements)

            for idx, (start, end) in enumerate(lines):
                count = elements_per_line + (1 if idx < extra else 0)
                if count == 0:
                    continue

                dx = end['x'] - start['x']
                dy = end['y'] - start['y']

                if dx == 0 and dy != 0:
                    step = (abs(dy) / count) / 2
                    pos_y = min(start['y'], end['y']) + step
                    for _ in range(count):
                        pos = {'x': start['x'], 'y': pos_y}
                        el = next(element_iter)
                        if el == 'resistor':
                            draw_resistor(pos, f'R{resistor_idx}', 'vertical')
                            resistor_idx += 1
                        elif el == 'capacitor':
                            draw_capacitor(pos, f'C{capacitor_idx}', 'vertical')
                            capacitor_idx += 1
                        elif el == 'inductor':
                            draw_inductor(pos, f'L{inductor_idx}', 'vertical')
                            inductor_idx += 1
                        pos_y += 2 * step

                elif dy == 0 and dx != 0:
                    step = (abs(dx) / count) / 2
                    pos_x = min(start['x'], end['x']) + step
                    for _ in range(count):
                        pos = {'x': pos_x, 'y': start['y']}
                        el = next(element_iter)
                        if el == 'resistor':
                            draw_resistor(pos, f'R{resistor_idx}', 'horizontal')
                            resistor_idx += 1
                        elif el == 'capacitor':
                            draw_capacitor(pos, f'C{capacitor_idx}', 'horizontal')
                            capacitor_idx += 1
                        elif el == 'inductor':
                            draw_inductor(pos, f'L{inductor_idx}', 'horizontal')
                            inductor_idx += 1
                        pos_x += 2 * step

                else:
                    length = np.sqrt(dx ** 2 + dy ** 2)
                    angle = np.arctan2(dy, dx)

                    if abs(angle) < np.pi / 2:
                        orientation = 'diagonal' if dy > 0 else '-diagonal'
                    else:
                        orientation = '-diagonal' if dy > 0 else 'diagonal'

                    valid_t = np.linspace(0.1, 0.9, count * 2)
                    valid_t = [t for t in valid_t if t < 0.4 or t > 0.6][:count]

                    for t in valid_t:
                        pos_x = start['x'] + t * dx
                        pos_y = start['y'] + t * dy
                        pos = {'x': pos_x, 'y': pos_y}
                        el = next(element_iter)
                        if el == 'resistor':
                            draw_resistor(pos, f'R{resistor_idx}', orientation)
                            resistor_idx += 1
                        elif el == 'capacitor':
                            draw_capacitor(pos, f'C{capacitor_idx}', orientation)
                            capacitor_idx += 1
                        elif el == 'inductor':
                            draw_inductor(pos, f'L{inductor_idx}', orientation)
                            inductor_idx += 1

    plt.axis('off')


def add_schemes_to_word(resistors_num, inductors_num, capacitors_num):
    doc = Document()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.allow_autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер варианта'
    hdr_cells[1].text = 'Схема'
    hdr_cells[2].text = 'Номиналы'
    hdr_cells[0].width = NUMBER_COL_WIDTH
    hdr_cells[1].width = MAX_IMAGE_WIDTH_INCHES
    hdr_cells[2].width = Inches(2.5)

    for i in range(1, 31):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[0].width = NUMBER_COL_WIDTH
        row_cells[1].width = MAX_IMAGE_WIDTH_INCHES
        row_cells[2].width = Inches(2.5)

        image_path = os.path.join(SCHEMES_FOLDER, f'scheme_{i}.png')
        if os.path.exists(image_path):
            with Image.open(image_path) as img:
                width_px, height_px = img.size
                dpi = img.info.get('dpi', (300, 300))[0]
                width_in = width_px / dpi
                scale = min(1.0, MAX_IMAGE_WIDTH_INCHES / width_in)
                display_width = Inches(width_in * scale)
                row_cells[1].paragraphs[0].add_run().add_picture(image_path, width=display_width)
        else:
            row_cells[1].text = 'Изображение не найдено'

        descriptions = []

        if resistors_num > 0:
            for r in range(1, resistors_num + 1):
                resistance = random.randint(5, 100)
                descriptions.append(f"R{r}={resistance} Ом")

        if capacitors_num > 0:
            for r in range(1, capacitors_num + 1):
                capacity = round(random.uniform(0.05, 1), 2)
                descriptions.append(f"C{r}={capacity} мкФ")

        if inductors_num > 0:
            for r in range(1, inductors_num + 1):
                inductance = random.randint(1, 20)
                descriptions.append(f"L{r}={inductance} мГн")

        row_cells[2].text = '\n'.join(descriptions)

    doc.save(OUTPUT_DOCX)
    print(f'Файл сохранён: {OUTPUT_DOCX}')


def generate_filter_schemes_set(scheme_type, filter_type):
    schemes = []
    resistors_num = 0
    capacitors_num = 0
    inductors_num = 0

    if filter_type == 'LPF':
        if scheme_type == 'G':
            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('resistor')
                scheme_layout['node2->node5'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                capacitors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('inductor')
                scheme_layout['node2->node5'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
            resistors_num = 1
            inductors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('inductor')
                scheme_layout['node2->node5'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                capacitors_num = 1
                inductors_num = 1

        if scheme_type == 'P':
            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('resistor')
                scheme_layout['node2->node6'][0]['elements'].append('capacitor')
                scheme_layout['node3->node7'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                capacitors_num = 2

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('inductor')
                scheme_layout['node2->node6'][0]['elements'].append('resistor')
                scheme_layout['node3->node7'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 2
                inductors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('inductor')
                scheme_layout['node2->node6'][0]['elements'].append('capacitor')
                scheme_layout['node3->node7'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                capacitors_num = 2
                inductors_num = 1

        if scheme_type == 'T':
            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('resistor')
                scheme_layout['node2->node3'][0]['elements'].append('resistor')
                scheme_layout['node2->node5'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 2
                capacitors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('inductor')
                scheme_layout['node2->node3'][0]['elements'].append('inductor')
                scheme_layout['node2->node5'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                inductors_num = 2

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('inductor')
                scheme_layout['node2->node3'][0]['elements'].append('inductor')
                scheme_layout['node2->node5'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                capacitors_num = 1
                inductors_num = 2

    if filter_type == 'HPF':
        if scheme_type == 'G':
            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('capacitor')
                scheme_layout['node2->node5'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                capacitors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('resistor')
                scheme_layout['node2->node5'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                inductors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('capacitor')
                scheme_layout['node2->node5'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                capacitors_num = 1
                inductors_num = 1

        if scheme_type == 'P':
            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('capacitor')
                scheme_layout['node2->node6'][0]['elements'].append('resistor')
                scheme_layout['node3->node7'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 2
                capacitors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('resistor')
                scheme_layout['node2->node6'][0]['elements'].append('inductor')
                scheme_layout['node3->node7'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                inductors_num = 2

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('capacitor')
                scheme_layout['node2->node6'][0]['elements'].append('inductor')
                scheme_layout['node3->node7'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                capacitors_num = 1
                inductors_num = 2

        if scheme_type == 'T':
            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('capacitor')
                scheme_layout['node2->node3'][0]['elements'].append('capacitor')
                scheme_layout['node2->node5'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                capacitors_num = 2

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('resistor')
                scheme_layout['node2->node3'][0]['elements'].append('resistor')
                scheme_layout['node2->node5'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 2
                inductors_num = 1

            for _ in range(11):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('capacitor')
                scheme_layout['node2->node3'][0]['elements'].append('capacitor')
                scheme_layout['node2->node5'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                capacitors_num = 2
                inductors_num = 1

    if filter_type == 'BPF':
        if scheme_type == 'G':
            for _ in range(31):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('inductor')
                scheme_layout['node1->node2'][0]['elements'].append('capacitor')
                scheme_layout['node2->node5'][0]['elements'].append('resistor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                capacitors_num = 1
                inductors_num = 1

        if scheme_type == 'P':
            for _ in range(31):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('capacitor')
                scheme_layout['node2->node3'][0]['elements'].append('resistor')
                scheme_layout['node2->node6'][0]['elements'].append('resistor')
                scheme_layout['node3->node7'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 2
                capacitors_num = 2

        if scheme_type == 'T':
            for _ in range(31):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('resistor')
                scheme_layout['node2->node3'][0]['elements'].append('inductor')
                scheme_layout['node2->node5'][0]['elements'].append('inductor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                inductors_num = 2

    if filter_type == 'BSF':
        if scheme_type == 'G':
            for _ in range(31):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node1->node2'][0]['elements'].append('resistor')
                scheme_layout['node2->node5'][0]['elements'].append('inductor')
                scheme_layout['node2->node5'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 1
                capacitors_num = 1
                inductors_num = 1

        if scheme_type == 'T_back_coupling':
            for _ in range(31):
                scheme_nodes, quadripole_nodes, scheme_layout = generate_filter_scheme_topology(
                    scheme_type=scheme_type
                )
                scheme_layout['node2->node3'][0]['elements'].append('capacitor')
                scheme_layout['node3->node5'][0]['elements'].append('capacitor')
                scheme_layout['node2->node4'][0]['elements'].append('resistor')
                scheme_layout['node4->node5'][0]['elements'].append('resistor')
                scheme_layout['node3->node8'][0]['elements'].append('resistor')
                scheme_layout['node4->node9'][0]['elements'].append('capacitor')
                schemes.append({
                    "scheme_nodes": scheme_nodes,
                    "quadripole_nodes": quadripole_nodes,
                    "scheme_layout": scheme_layout
                })
                resistors_num = 3
                capacitors_num = 3

    unique_schemes = []

    for scheme in schemes:
        if len(unique_schemes) == 30:
            break
        if scheme not in unique_schemes:
            unique_schemes.append(scheme)

    if len(unique_schemes) < 30:
        print(f'[Warning] Удалось сгенерировать только {len(unique_schemes)} уникальных схем')
        while len(unique_schemes) < 30:
            random_scheme = random.choice(schemes)
            unique_schemes.append(random_scheme)

    random.shuffle(unique_schemes)

    for idx, scheme in enumerate(unique_schemes, start=1):
        visualise_filter_scheme(
            scheme_nodes=scheme["scheme_nodes"],
            quadripole_nodes=scheme["quadripole_nodes"],
            scheme_layout=scheme["scheme_layout"]
        )

        fig = plt.gcf()
        fig.savefig(f'{SCHEMES_FOLDER}/scheme_{idx}.png', dpi=300, bbox_inches='tight', pad_inches=0)
        plt.close(fig)

    add_schemes_to_word(
        resistors_num=resistors_num,
        capacitors_num=capacitors_num,
        inductors_num=inductors_num
    )
