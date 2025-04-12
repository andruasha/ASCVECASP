import random
import matplotlib.pyplot as plt
import os
import itertools

from docx import Document
from docx.shared import Inches
from PIL import Image

from src.alternating_current.electric_circuit import ElectricCircuit
from src.alternating_current.elements_places import ElementsPlacer
from src.alternating_current.visualize_circuit import CircuitVisualize

from conf.config import SCALE


MAX_IMAGE_WIDTH_INCHES = 3
NUMBER_COL_WIDTH = Inches(0.8)
SCHEMES_FOLDER = 'schemes'
OUTPUT_DOCX = 'generated_schemes.docx'


def compare_topologies(circuit_topology_1, circuit_topology_2):
    def normalize_connection(node_a, node_b):
        return "->".join(sorted([node_a, node_b]))

    def get_topology_representation(circuit_topology):
        representation = {}
        for conn_name, conn_coords in circuit_topology.nodes_connections.items():
            node1, node2 = conn_name.split('->')
            key = normalize_connection(node1, node2)
            if key not in representation:
                representation[key] = 0
            representation[key] += len(conn_coords)
        return representation

    rep1 = get_topology_representation(circuit_topology_1)
    rep2 = get_topology_representation(circuit_topology_2)

    return rep1 == rep2


def compare_layouts(circuit_layout_1, circuit_layout_2):
    absolutely_unique = True
    for connection_name_1, connection_params_1 in circuit_layout_1.items():
        for connection_name_2, connection_params_2 in circuit_layout_2.items():
            splited_name_1 = connection_name_1.split('->')
            if (connection_name_2 == connection_params_1) or (
                    connection_name_2 == f"{splited_name_1[1]}->{splited_name_1[0]}"):
                if len(connection_params_1) == len(connection_params_2):
                    connection_elements_1 = []
                    connection_elements_2 = []

                    for sub_connection in connection_params_1:
                        connection_elements_1.append(sub_connection[0]['elements'])

                    for sub_connection in connection_params_2:
                        connection_elements_2.append(sub_connection[0]['elements'])

                    set1 = {tuple(sorted(tuple(element.items()) for element in sub_list))
                            for sub_list in connection_elements_1}
                    set2 = {tuple(sorted(tuple(element.items()) for element in sub_list))
                            for sub_list in connection_elements_2}

                    if set1 != set2:
                        absolutely_unique = False

    return absolutely_unique


def generate_schemes_set(nodes_num, branches_num, voltage_sources_num, current_sources_num, resistors_num, inductors_num, capacitors_num):
    template_two_nodes_1 = {'node1': {'x': 0, 'y': 0},
                            'node2': {'x': SCALE, 'y': 0}}

    template_three_nodes_1 = {'node1': {'x': 0, 'y': 0},
                              'node2': {'x': 0, 'y': SCALE},
                              'node3': {'x': SCALE, 'y': SCALE}}

    template_four_nodes_1 = {'node1': {'x': 0, 'y': 0},
                             'node2': {'x': 0, 'y': SCALE},
                             'node3': {'x': SCALE, 'y': SCALE},
                             'node4': {'x': SCALE, 'y': 0}}

    template_four_nodes_2 = {'node1': {'x': 0, 'y': SCALE},
                             'node2': {'x': SCALE, 'y': SCALE},
                             'node3': {'x': 2 * SCALE, 'y': SCALE},
                             'node4': {'x': SCALE, 'y': 0}}

    def get_unique_topologies(topologies, needed_count):
        unique = []
        for topo in topologies:
            if all(not compare_topologies(topo, u) for u in unique):
                unique.append(topo)
                if len(unique) >= needed_count:
                    break
        return unique

    circuits_topologies = []

    if nodes_num == 2:
        raw_topologies = [ElectricCircuit(branches_num=branches_num, nodes=template_two_nodes_1).create_nodes_connections() for _ in range(100)]
        circuits_topologies = get_unique_topologies(raw_topologies, 30)

    elif nodes_num == 3:
        raw_topologies = [ElectricCircuit(branches_num=branches_num, nodes=template_three_nodes_1).create_nodes_connections() for _ in range(500)]
        circuits_topologies = get_unique_topologies(raw_topologies, 30)

    elif nodes_num == 4:
        circuits_topologies_1 = [ElectricCircuit(branches_num=branches_num, nodes=template_four_nodes_1).create_nodes_connections() for _ in range(300)]
        circuits_topologies_2 = [ElectricCircuit(branches_num=branches_num, nodes=template_four_nodes_2).create_nodes_connections() for _ in range(300)]

        unique_1 = get_unique_topologies(circuits_topologies_1, 15)
        unique_2 = get_unique_topologies(circuits_topologies_2, 15)
        circuits_topologies = unique_1 + unique_2

        if len(circuits_topologies) < 30:
            print(f'[Warning] Удалось сгенерировать только {len(circuits_topologies)} уникальных топологий.')
            print(f'[Info] Добираем ещё {30 - len(circuits_topologies)} топологий случайным образом (с возможными повторами).')
            all_topos = circuits_topologies_1 + circuits_topologies_2
            while len(circuits_topologies) < 30:
                circuits_topologies.append(random.choice(all_topos))

        random.shuffle(circuits_topologies)

    if len(circuits_topologies) < 30:
        print(f'[Warning] Удалось сгенерировать только {len(circuits_topologies)} уникальных топологий.')
        print(f'[Info] Добираем ещё {30 - len(circuits_topologies)} топологий случайным образом (с возможными повторами).')
        while len(circuits_topologies) < 30:
            circuits_topologies.append(random.choice(raw_topologies))

    circuits = []
    circuits_topologies_paired = []

    for circuit_topology in circuits_topologies:
        placed_circuit = ElementsPlacer(circuit_topology,
                                        voltage_sources_num,
                                        current_sources_num,
                                        resistors_num,
                                        inductors_num,
                                        capacitors_num).place_elements()
        circuits.append(placed_circuit)
        circuits_topologies_paired.append((circuit_topology, placed_circuit))

    unique_schemes = []

    for i in range(len(circuits_topologies_paired)):
        topology_1, circuit_1 = circuits_topologies_paired[i]
        is_unique = True

        for existing_topology, existing_circuit in unique_schemes:
            if compare_layouts(circuit_1.layout, existing_circuit.layout) and compare_topologies(topology_1, existing_topology):
                is_unique = False
                break

        if is_unique:
            unique_schemes.append((topology_1, circuit_1))

    if len(unique_schemes) < 30:
        print(f'[Warning] Удалось сгенерировать только {len(unique_schemes)} уникальных схем')

    index = 1
    for topology, circuit in circuits_topologies_paired:
        visualizer = CircuitVisualize(circuit, topology)
        visualizer.visualize()

        fig = plt.gcf()
        fig.savefig(f'{SCHEMES_FOLDER}/scheme_{index}.png', dpi=300, bbox_inches='tight', pad_inches=0)
        plt.close(fig)

        # ltspice_path = os.path.join(SCHEMES_FOLDER, f'scheme_{index}.cir')
        # visualizer.export_to_ltspice(filename=ltspice_path)

        index += 1


def add_schemes_to_word(voltage_sources_num, current_sources_num, resistors_num, inductors_num, capacitors_num):
    doc = Document()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.allow_autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер варианта'
    hdr_cells[1].text = 'Схема'
    hdr_cells[2].text = 'Номиналы'
    hdr_cells[0].width = NUMBER_COL_WIDTH
    hdr_cells[1].width = MAX_IMAGE_WIDTH_INCHES
    hdr_cells[2].width = Inches(2.5)

    for i in range(1, 31):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[0].width = NUMBER_COL_WIDTH
        row_cells[1].width = MAX_IMAGE_WIDTH_INCHES
        row_cells[2].width = Inches(2.5)

        image_path = os.path.join(SCHEMES_FOLDER, f'scheme_{i}.png')
        if os.path.exists(image_path):
            with Image.open(image_path) as img:
                width_px, height_px = img.size
                dpi = img.info.get('dpi', (300, 300))[0]
                width_in = width_px / dpi
                scale = min(1.0, MAX_IMAGE_WIDTH_INCHES / width_in)
                display_width = Inches(width_in * scale)
                row_cells[1].paragraphs[0].add_run().add_picture(image_path, width=display_width)
        else:
            row_cells[1].text = 'Изображение не найдено'

        descriptions = []

        for v in range(1, voltage_sources_num + 1):
            voltage = random.randint(15, 310)
            descriptions.append(f"V{v}={voltage} В")

        for r in range(1, resistors_num + 1):
            resistance = random.randint(5, 100)
            descriptions.append(f"R{r}={resistance} Ом")

        for r in range(1, capacitors_num + 1):
            capacity = round(random.uniform(0.05, 1), 2)
            descriptions.append(f"C{r}={capacity} мкФ")

        for r in range(1, inductors_num + 1):
            inductance = random.randint(1, 20)
            descriptions.append(f"L{r}={inductance} мГн")

        for c in range(1, current_sources_num + 1):
            current = round(random.uniform(0.15, 3), 2)
            descriptions.append(f"I{c}={current} А")

        frequency = random.randint(1, 20)
        descriptions.append(f"f={frequency} кГц")

        row_cells[2].text = '\n'.join(descriptions)

    doc.save(OUTPUT_DOCX)
    print(f'Файл сохранён: {OUTPUT_DOCX}')


def generate_alternating_current_schemes_set(nodes_num, branches_num, voltage_sources_num, current_sources_num, resistors_num, inductors_num, capacitors_num):
    generate_schemes_set(
        nodes_num=nodes_num,
        branches_num=branches_num,
        voltage_sources_num=voltage_sources_num,
        current_sources_num=current_sources_num,
        resistors_num=resistors_num,
        capacitors_num=capacitors_num,
        inductors_num=inductors_num
    )

    add_schemes_to_word(
        voltage_sources_num=voltage_sources_num,
        current_sources_num=current_sources_num,
        resistors_num=resistors_num,
        capacitors_num=capacitors_num,
        inductors_num=inductors_num
    )
