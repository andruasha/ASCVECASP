import random
import matplotlib.pyplot as plt
import os
import itertools

from docx import Document
from docx.shared import Inches
from PIL import Image

from src.electric_circuit import ElectricCircuit
from src.elements_places import ElementsPlacer
from src.visualize_circuit import CircuitVisualize

from conf.config import SCALE


MAX_IMAGE_WIDTH_INCHES = 3
NUMBER_COL_WIDTH = Inches(0.8)
SCHEMES_FOLDER = 'schemes'
OUTPUT_DOCX = 'generated_schemes.docx'


def compare_topologies(circuit_topology_1, circuit_topology_2):
    def normalize_connection(node_a, node_b):
        return "->".join(sorted([node_a, node_b]))

    def get_topology_representation(circuit_topology):
        representation = {}
        for conn_name, conn_coords in circuit_topology.nodes_connections.items():
            node1, node2 = conn_name.split('->')
            key = normalize_connection(node1, node2)
            if key not in representation:
                representation[key] = 0
            representation[key] += len(conn_coords)
        return representation

    rep1 = get_topology_representation(circuit_topology_1)
    rep2 = get_topology_representation(circuit_topology_2)

    return rep1 == rep2


def generate_schemes_set(nodes_num, branches_num, voltage_sources_num, current_sources_num, resistors_num, inductors_num,
                         capacitors_num):
    template_two_nodes_1 = {'node1': {'x': 0, 'y': 0},
                            'node2': {'x': SCALE, 'y': 0}}

    template_three_nodes_1 = {'node1': {'x': 0, 'y': 0},
                              'node2': {'x': 0, 'y': SCALE},
                              'node3': {'x': SCALE, 'y': SCALE}}

    template_four_nodes_1 = {'node1': {'x': 0, 'y': 0},
                             'node2': {'x': 0, 'y': SCALE},
                             'node3': {'x': SCALE, 'y': SCALE},
                             'node4': {'x': SCALE, 'y': 0}}

    template_four_nodes_2 = {'node1': {'x': 0, 'y': SCALE},
                             'node2': {'x': SCALE, 'y': SCALE},
                             'node3': {'x': 2 * SCALE, 'y': SCALE},
                             'node4': {'x': SCALE, 'y': 0}}

    circuits_topologies = []

    if nodes_num == 2:
        circuits_topologies = [ElectricCircuit(branches_num=branches_num, nodes=template_two_nodes_1).create_nodes_connections() for _ in range(5000)]
    elif nodes_num == 3:
        circuits_topologies = [ElectricCircuit(branches_num=branches_num, nodes=template_three_nodes_1).create_nodes_connections() for _ in range(10000)]
    elif nodes_num == 4:
        circuits_topologies = [ElectricCircuit(branches_num=branches_num, nodes=template_four_nodes_1).create_nodes_connections() for _ in range(15)]
        circuits_topologies += [ElectricCircuit(branches_num=branches_num, nodes=template_four_nodes_2).create_nodes_connections() for _ in range(15)]
        random.shuffle(circuits_topologies)

    unique_circuits_topologies = []
    for circuit_topology in circuits_topologies:
        is_unique = True
        for unique_circuit in list(unique_circuits_topologies):
            if compare_topologies(circuit_topology, unique_circuit):
                is_unique = False
                break
        if is_unique:
            unique_circuits_topologies.append(circuit_topology)
            if len(unique_circuits_topologies) >= 30:
                break

    if len(unique_circuits_topologies) < 30:
        missing = 30 - len(unique_circuits_topologies)
        print(f'[Warning] Удалось сгенерировать только {len(unique_circuits_topologies)} уникальных топологий.')
        print(f'[Info] Добираем ещё {missing} топологий случайным образом (с возможными повторами).')
        while len(unique_circuits_topologies) < 30:
            random_topology = random.choice(circuits_topologies)
            unique_circuits_topologies.append(random_topology)

    circuits = []
    circuits_topologies_paired = []

    for circuit_topology in unique_circuits_topologies:
        placed_circuit = ElementsPlacer(circuit_topology,
                                        voltage_sources_num,
                                        current_sources_num,
                                        resistors_num,
                                        inductors_num,
                                        capacitors_num).place_elements()
        circuits.append(placed_circuit)
        circuits_topologies_paired.append((circuit_topology, placed_circuit))

    index = 1
    for topology, circuit in circuits_topologies_paired:
        CircuitVisualize(circuit, topology).visualize()
        fig = plt.gcf()
        fig.savefig(f'{SCHEMES_FOLDER}/scheme_{index}.png', dpi=300, bbox_inches='tight', pad_inches=0)
        plt.close(fig)
        index += 1


def add_schemes_to_word():
    doc = Document()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.allow_autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер варианта'
    hdr_cells[1].text = 'Схема'
    hdr_cells[0].width = NUMBER_COL_WIDTH
    hdr_cells[1].width = MAX_IMAGE_WIDTH_INCHES

    for i in range(1, 31):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[0].width = NUMBER_COL_WIDTH
        row_cells[1].width = MAX_IMAGE_WIDTH_INCHES

        image_path = os.path.join(SCHEMES_FOLDER, f'scheme_{i}.png')
        if os.path.exists(image_path):
            with Image.open(image_path) as img:
                width_px, height_px = img.size
                dpi = img.info.get('dpi', (300, 300))[0]
                width_in = width_px / dpi
                scale = min(1.0, MAX_IMAGE_WIDTH_INCHES / width_in)
                display_width = Inches(width_in * scale)
                row_cells[1].paragraphs[0].add_run().add_picture(image_path, width=display_width)
        else:
            row_cells[1].text = 'Изображение не найдено'

    doc.save(OUTPUT_DOCX)
    print(f'Файл сохранён: {OUTPUT_DOCX}')


generate_schemes_set(
    nodes_num=4,
    branches_num=10,
    voltage_sources_num=1,
    current_sources_num=2,
    resistors_num=4,
    inductors_num=2,
    capacitors_num=2
)

add_schemes_to_word()
